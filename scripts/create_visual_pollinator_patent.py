from pathlib import Path
import math

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont, ImageOps
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as RLImage,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "patent"
ASSETS = OUT / "visual_assets"
OUT.mkdir(parents=True, exist_ok=True)
ASSETS.mkdir(parents=True, exist_ok=True)

TITLE = "Weather-Adaptive Pollinator Activity Prediction Using Ensemble Learning and Real-Time Meteorological Streams"
DOCX_PATH = OUT / "Weather_Adaptive_Pollinator_Activity_Prediction_VISUAL_PATENT.docx"
PDF_PATH = OUT / "Weather_Adaptive_Pollinator_Activity_Prediction_VISUAL_PATENT.pdf"

A_SIGNATURE = Path(r"C:\Users\aarav\Downloads\WhatsApp Image 2026-05-08 at 8.39.12 PM.jpeg")
B_SIGNATURE = Path(r"C:\Users\aarav\Downloads\WhatsApp Image 2026-05-08 at 8.58.33 PM.jpeg")

inventors = [
    {
        "label": "A",
        "name": "Aarav Kashyap Singh",
        "mobile": "6281328903",
        "email": "aaravkashyap1203@gmail.com",
        "reg": "12305324",
        "address": "Lovely Professional University, Punjab-144411, India",
        "signature": A_SIGNATURE,
    },
    {
        "label": "B",
        "name": "Jai Dev Meena",
        "mobile": "8209926739",
        "email": "mjai9127@gmail.com",
        "reg": "12318662",
        "address": "Lovely Professional University, Punjab-144411, India",
        "signature": B_SIGNATURE,
    },
]

BLUE = "#1F4E79"
NAVY = "#16324F"
CYAN = "#DDEFF8"
GREEN = "#DCEEDC"
YELLOW = "#FFF2CC"
ORANGE = "#FCE4D6"
PURPLE = "#EADCF8"
RED = "#F4CCCC"
GRAY = "#F4F6F8"
W, H = 1600, 940
HEADER_H = 96
PANEL = (60, 140, 1540, 870)


def font(size, bold=False):
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


F_TITLE = font(34, True)
F_H = font(22, True)
F_B = font(18, False)
F_SB = font(16, True)
F_S = font(14, False)
F_XS = font(12, False)


def wrap_text(draw, text, max_width, fnt):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        trial = (current + " " + word).strip()
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def rounded(draw, xy, fill, outline="#2F3A45", width=2, radius=18):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def centered_text(draw, box, text, fnt, fill="#111111", spacing=4):
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, x2 - x1 - 20, fnt)
    line_h = fnt.size + spacing
    total_h = len(lines) * line_h
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        w = draw.textbbox((0, 0), line, font=fnt)[2]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill=fill)
        y += line_h


def arrow(draw, start, end, fill="#263238", width=5):
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 15
    p1 = (end[0] - size * math.cos(angle - math.pi / 6), end[1] - size * math.sin(angle - math.pi / 6))
    p2 = (end[0] - size * math.cos(angle + math.pi / 6), end[1] - size * math.sin(angle + math.pi / 6))
    draw.polygon([end, p1, p2], fill=fill)


def save_diagram(name, title, draw_body):
    img = Image.new("RGB", (W, H), "#FBFCFE")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, W, H), fill="#FBFCFE")
    for x in range(0, W, 40):
        draw.line((x, HEADER_H, x, H), fill="#ECF2F7", width=1)
    for y in range(HEADER_H, H, 40):
        draw.line((0, y, W, y), fill="#ECF2F7", width=1)
    draw.rectangle((0, 0, W, HEADER_H), fill=BLUE)
    draw.rectangle((0, HEADER_H, W, HEADER_H + 9), fill="#8CC63F")
    draw.text((40, 26), title, font=F_TITLE, fill="white")
    draw.rounded_rectangle(PANEL, radius=28, outline="#8CAAC0", width=4, fill="#FFFFFF")
    draw_body(draw)
    draw.rectangle((60, 890, 1540, 910), fill="#EAF2F8", outline="#8CAAC0", width=1)
    out = ASSETS / name
    img.save(out, quality=95)
    return out


def create_architecture():
    def body(d):
        xs = [105, 395, 685, 975, 1265]
        input_boxes = [
            ("Weather APIs\nFarm Weather Station", CYAN),
            ("IoT Sensors\nCamera/Acoustic Counters", GREEN),
            ("Crop Phenology\nBloom Stage and Density", YELLOW),
            ("Landscape Context\nHabitat and Field Map", ORANGE),
            ("Manual Survey\nHive and Visit Logs", PURPLE),
        ]
        for x, (txt, col) in zip(xs, input_boxes):
            rounded(d, (x, 170, x + 230, 260), col)
            centered_text(d, (x, 170, x + 230, 260), txt, F_S)
            arrow(d, (x + 115, 260), (800, 345), width=4)
        rounded(d, (430, 345, 1170, 430), "#E2F0D9", outline=NAVY, width=4)
        centered_text(d, (430, 345, 1170, 430), "Stream Synchronization + Weather-Aware Feature Engineering", F_H, fill=NAVY)
        for i, label in enumerate(["Lag", "VPD", "Rain Recovery", "Wind Shock", "Solar Phase", "Bloom Demand"]):
            x = 475 + i * 110
            rounded(d, (x, 455, x + 98, 495), "#FFFFFF", outline="#6C8EBF", width=2, radius=12)
            centered_text(d, (x, 455, x + 98, 495), label, F_XS, fill=NAVY)
        model_boxes = [
            ("Gradient\nBoosting", "#D9EAD3"),
            ("Random\nForest", "#D0E0E3"),
            ("Temporal\nNeural Net", "#D9D2E9"),
            ("Bayesian / Quantile\nUncertainty", "#FCE5CD"),
            ("Logistic\nActivity Classifier", "#F4CCCC"),
        ]
        for x, (txt, col) in zip(xs, model_boxes):
            rounded(d, (x, 545, x + 230, 635), col)
            centered_text(d, (x, 545, x + 230, 635), txt, F_SB)
            arrow(d, (800, 430), (x + 115, 545), width=4)
            arrow(d, (x + 115, 635), (800, 710), width=4)
        rounded(d, (430, 710, 1170, 795), "#CFE2F3", outline=NAVY, width=4)
        centered_text(d, (430, 710, 1170, 795), "Stacking Meta-Learner + Calibration + Drift Detection", F_H, fill=NAVY)
        outputs = [
            ("Activity Index", "#D9EAD3"),
            ("Peak Window", "#FFF2CC"),
            ("Deficit Risk", "#F4CCCC"),
            ("Pesticide Caution", "#FCE4D6"),
            ("Dashboard + API", "#D9D2E9"),
        ]
        for x, (txt, col) in zip(xs, outputs):
            rounded(d, (x, 815, x + 230, 865), col)
            centered_text(d, (x, 815, x + 230, 865), txt, F_XS)
            arrow(d, (800, 795), (x + 115, 815), width=3)
    return save_diagram("architecture.png", "FIG. 1 - Complete Weather-Adaptive Pollinator Prediction Architecture", body)


def create_flow():
    def body(d):
        steps = [
            ("1. Collect live weather + field observations", CYAN),
            ("2. Validate data quality, fill gaps, align timestamps", GREEN),
            ("3. Build rolling weather, crop, and activity features", YELLOW),
            ("4. Run parallel ensemble models", PURPLE),
            ("5. Combine model outputs through meta-learner", "#CFE2F3"),
            ("6. Calibrate probability and compute confidence", ORANGE),
            ("7. Generate forecast, risk alert, and recommendation", RED),
            ("8. Store feedback for continuous learning", "#D9EAD3"),
        ]
        y = 165
        for i, (txt, col) in enumerate(steps):
            x = 160 if i % 2 == 0 else 890
            rounded(d, (x, y, x + 550, y + 72), col, outline=NAVY, width=3)
            d.rectangle((x, y, x + 58, y + 72), fill=NAVY)
            centered_text(d, (x, y, x + 58, y + 72), str(i + 1), F_H, fill="white")
            centered_text(d, (x + 60, y, x + 550, y + 72), txt, F_SB, fill=NAVY)
            if i < len(steps) - 1:
                next_x = 180 if (i + 1) % 2 == 0 else 870
                next_x = 160 if (i + 1) % 2 == 0 else 890
                next_y = y + 82
                arrow(d, (x + 550, y + 36) if next_x > x else (x, y + 36), (next_x, next_y + 36) if next_x > x else (next_x + 550, next_y + 36), width=5)
            y += 82
        d.rectangle((150, 825, 1450, 860), fill="#EEF5FB", outline=NAVY, width=2)
        centered_text(d, (150, 825, 1450, 860), "Refresh cycle: 5, 15, 30, or 60 minutes depending on sensor availability.", F_S, fill=NAVY)
    return save_diagram("prediction_flowchart.png", "FIG. 2 - Prediction Pipeline Flowchart", body)


def create_training_loop():
    def body(d):
        center = (800, 540)
        circle_items = [
            ("Historical Data", 800, 185, CYAN),
            ("Feature Store", 1180, 320, GREEN),
            ("Train Base Models", 1180, 635, YELLOW),
            ("Validate + Calibrate", 800, 795, ORANGE),
            ("Deploy Ensemble", 420, 635, PURPLE),
            ("Live Feedback", 420, 320, RED),
        ]
        prev = None
        first = None
        centers = []
        for txt, x, y, col in circle_items:
            rounded(d, (x - 150, y - 55, x + 150, y + 55), col, outline=NAVY, width=3)
            centered_text(d, (x - 150, y - 55, x + 150, y + 55), txt, F_SB, fill=NAVY)
            centers.append((x, y))
            if first is None:
                first = (x, y)
            if prev:
                arrow(d, prev, (x, y), width=4)
            prev = (x, y)
        arrow(d, prev, first, width=4)
        rounded(d, (570, 430, 1030, 625), "#EAF2F8", outline=NAVY, width=5, radius=28)
        centered_text(d, (570, 430, 1030, 625), "Dynamic Ensemble Weighting\nModel Drift Detection\nWeather Regime Adaptation", F_H, fill=NAVY)
        for i, pct in enumerate([82, 76, 89, 71]):
            x = 620 + i * 100
            d.rectangle((x, 650, x + 52, 790), outline=NAVY, width=2, fill="#EEF5FB")
            d.rectangle((x, 790 - pct, x + 52, 790), fill=["#6AA84F", "#3D85C6", "#674EA7", "#E69138"][i])
            centered_text(d, (x - 5, 798, x + 58, 830), f"{pct}%", F_XS, fill=NAVY)
        for c in centers:
            arrow(d, c, center, fill="#7F8C8D", width=3)
    return save_diagram("training_loop.png", "FIG. 3 - Continuous Learning and Model Recalibration Loop", body)


def create_dashboard():
    def body(d):
        d.text((105, 165), "Farm Dashboard - Pollinator Activity Forecast", font=F_H, fill=NAVY)
        for i, label in enumerate(["Farm A", "Crop: Mustard", "Horizon: 6h", "Confidence: 91%"]):
            rounded(d, (780 + i * 175, 155, 935 + i * 175, 195), "#FFFFFF", outline="#6C8EBF", width=2, radius=12)
            centered_text(d, (780 + i * 175, 155, 935 + i * 175, 195), label, F_XS, fill=NAVY)
        rounded(d, (110, 225, 530, 420), "#E2F0D9", outline="#6AA84F", width=3)
        centered_text(d, (110, 225, 530, 275), "Activity Index", F_H, fill="#274E13")
        d.text((225, 295), "0.82", font=font(66, True), fill="#38761D")
        centered_text(d, (110, 370, 530, 410), "High predicted activity from 09:30 to 11:45", F_S, fill="#274E13")
        rounded(d, (585, 225, 1015, 420), "#FFF2CC", outline="#BF9000", width=3)
        centered_text(d, (585, 225, 1015, 275), "Weather Drivers", F_H, fill="#7F6000")
        drivers = ["Temperature: favorable", "Wind gusts: moderate", "Humidity: stable", "Rain recovery: complete"]
        for i, t in enumerate(drivers):
            d.text((640, 300 + i * 25), t, font=F_S, fill="#7F6000")
        rounded(d, (1070, 225, 1490, 420), "#F4CCCC", outline="#CC0000", width=3)
        centered_text(d, (1070, 225, 1490, 275), "Operational Alert", F_H, fill="#990000")
        centered_text(d, (1110, 305, 1450, 380), "Avoid pesticide spray during active pollination window", F_S, fill="#990000")
        rounded(d, (110, 470, 800, 845), "#FFFFFF", outline=NAVY, width=3)
        d.text((140, 495), "Field Heat Map", font=F_SB, fill=NAVY)
        colors_map = ["#D9EAD3", "#B6D7A8", "#93C47D", "#FCE5CD", "#F4CCCC"]
        for row in range(5):
            for col in range(8):
                fill = colors_map[(row * 2 + col) % len(colors_map)]
                d.rounded_rectangle((150 + col * 75, 565 + row * 45, 210 + col * 75, 600 + row * 45), radius=8, fill=fill, outline="#777777")
        rounded(d, (860, 470, 1490, 845), "#FFFFFF", outline=NAVY, width=3)
        d.text((895, 495), "Hourly Forecast Timeline", font=F_SB, fill=NAVY)
        base_x, base_y = 930, 775
        d.line((base_x, base_y, 1430, base_y), fill="#555555", width=3)
        d.line((base_x, 585, base_x, base_y), fill="#555555", width=3)
        pts = [(930, 735), (1000, 700), (1070, 640), (1140, 610), (1210, 630), (1280, 690), (1350, 740), (1420, 760)]
        d.line(pts, fill="#1F4E79", width=7)
        for x, y in pts:
            d.ellipse((x - 7, y - 7, x + 7, y + 7), fill="#1F4E79")
        for i, lab in enumerate(["7", "8", "9", "10", "11", "12", "13", "14"]):
            d.text((920 + i * 70, 795), lab, font=F_XS, fill="#333333")
    return save_diagram("dashboard_mockup.png", "FIG. 4 - Decision Dashboard, Heat Map, and Alert Interface", body)


def create_risk_matrix():
    def body(d):
        d.text((95, 165), "Pollination deficit risk is computed by comparing predicted activity with crop-stage demand.", font=F_B, fill=NAVY)
        headers = ["Weather State", "Predicted Activity", "Crop Flowering Demand", "Risk Level", "Recommended Action"]
        rows = [
            ["Warm + calm + sunny", "High", "High", "Low", "Normal monitoring"],
            ["Windy + cloudy", "Medium", "High", "Medium", "Scout field / check hives"],
            ["Rain within 2 hours", "Low", "High", "High", "Delay pesticide / support pollination"],
            ["Heat stress afternoon", "Low", "Medium", "Medium", "Prioritize morning operations"],
            ["Cold morning recovery", "Rising", "High", "Medium", "Schedule activity after warm-up"],
        ]
        x0, y0 = 90, 230
        col_w = [265, 255, 300, 210, 400]
        row_h = 86
        x = x0
        for i, h in enumerate(headers):
            d.rectangle((x, y0, x + col_w[i], y0 + 70), fill=BLUE, outline="white", width=2)
            centered_text(d, (x, y0, x + col_w[i], y0 + 70), h, F_SB, fill="white")
            x += col_w[i]
        for r, row in enumerate(rows):
            x = x0
            y = y0 + 70 + r * row_h
            for i, cell in enumerate(row):
                fill = "#FFFFFF" if r % 2 == 0 else "#F4F6F8"
                if i == 3 and cell == "High":
                    fill = "#F4CCCC"
                elif i == 3 and cell == "Medium":
                    fill = "#FFF2CC"
                elif i == 3 and cell == "Low":
                    fill = "#D9EAD3"
                d.rectangle((x, y, x + col_w[i], y + row_h), fill=fill, outline="#777777", width=2)
                centered_text(d, (x, y, x + col_w[i], y + row_h), cell, F_B if i == 3 else F_S, fill=NAVY)
                x += col_w[i]
        rounded(d, (165, 800, 1435, 860), "#EAF2F8", outline=NAVY, width=4)
        centered_text(d, (165, 800, 1435, 860), "Risk Score = Crop Demand - Weather-Adjusted Activity + Uncertainty Penalty", F_H, fill=NAVY)
    return save_diagram("risk_matrix.png", "FIG. 5 - Pollination Deficit Risk Matrix and Action Logic", body)


def create_model_stack():
    def body(d):
        left = [
            ("Live Weather\nStream", CYAN),
            ("Pollinator\nDetections", GREEN),
            ("Crop Bloom\nDemand", YELLOW),
            ("Field Context\nMap", ORANGE),
        ]
        for i, (txt, col) in enumerate(left):
            y = 175 + i * 155
            rounded(d, (90, y, 360, y + 95), col, outline=NAVY, width=3)
            centered_text(d, (90, y, 360, y + 95), txt, F_SB, fill=NAVY)
            arrow(d, (360, y + 48), (500, 420), width=5)
        rounded(d, (500, 250, 755, 590), "#EAF2F8", outline=NAVY, width=5, radius=30)
        centered_text(d, (510, 260, 745, 335), "Feature Factory", F_H, fill=NAVY)
        chips = ["Rolling windows", "Lag values", "Microclimate stress", "Bloom demand", "Visit history", "Location encoding"]
        for i, chip in enumerate(chips):
            x = 535 + (i % 2) * 105
            y = 360 + (i // 2) * 62
            rounded(d, (x, y, x + 95, y + 42), "#FFFFFF", outline="#6C8EBF", width=2, radius=12)
            centered_text(d, (x, y, x + 95, y + 42), chip, F_XS, fill=NAVY)
        model_y = [165, 300, 435, 570, 705]
        models = [("GBM", "#D9EAD3"), ("RF", "#D0E0E3"), ("TNN", "#D9D2E9"), ("Bayes", "#FCE5CD"), ("Classifier", "#F4CCCC")]
        for (txt, col), y in zip(models, model_y):
            rounded(d, (865, y, 1105, y + 85), col, outline=NAVY, width=3)
            centered_text(d, (865, y, 1105, y + 85), txt, F_H, fill=NAVY)
            arrow(d, (755, 420), (865, y + 48), width=4)
            arrow(d, (1105, y + 48), (1235, 465), width=4)
        rounded(d, (1235, 345, 1515, 590), "#CFE2F3", outline=NAVY, width=5, radius=26)
        centered_text(d, (1245, 355, 1505, 450), "Meta-Learner", F_H, fill=NAVY)
        centered_text(d, (1260, 465, 1490, 560), "Weighted forecast\ncalibration\nuncertainty band", F_B, fill=NAVY)
        outputs = [("0.84 Activity", "#D9EAD3"), ("Low Risk", "#D9EAD3"), ("Peak 9-11 AM", "#FFF2CC")]
        for i, (txt, col) in enumerate(outputs):
            rounded(d, (1245 + i * 95, 690, 1330 + i * 95, 760), col, outline=NAVY, width=2, radius=14)
            centered_text(d, (1245 + i * 95, 690, 1330 + i * 95, 760), txt, F_XS, fill=NAVY)
            arrow(d, (1375, 590), (1288 + i * 95, 690), width=4)
    return save_diagram("model_stack.png", "FIG. 6 - Ensemble Model Stack and Confidence Engine", body)


def unified_node(draw, box, number, title, detail, fill):
    x1, y1, x2, y2 = box
    rounded(draw, box, fill, outline=NAVY, width=3, radius=18)
    draw.rectangle((x1, y1, x1 + 48, y2), fill=NAVY)
    centered_text(draw, (x1, y1, x1 + 48, y2), str(number), F_SB, fill="white")
    draw.text((x1 + 62, y1 + 14), title, font=F_SB, fill=NAVY)
    lines = wrap_text(draw, detail, x2 - x1 - 82, F_XS)
    y = y1 + 42
    for line in lines[:3]:
        draw.text((x1 + 62, y), line, font=F_XS, fill="#263238")
        y += 17


def create_unified_blueprint(name, fig_no, heading, subtitle, rows, metrics):
    def body(d):
        # Left identity rail
        d.rounded_rectangle((85, 165, 300, 845), radius=22, fill="#102A43", outline=NAVY, width=3)
        d.text((118, 205), f"FIG. {fig_no}", font=font(46, True), fill="white")
        d.line((115, 270, 270, 270), fill="#8CC63F", width=6)
        for i, line in enumerate(wrap_text(d, heading, 165, F_SB)[:5]):
            d.text((112, 300 + i * 24), line, font=F_SB, fill="white")
        d.text((112, 475), "FOCUS", font=F_XS, fill="#8CC63F")
        for i, line in enumerate(wrap_text(d, subtitle, 165, F_XS)[:7]):
            d.text((112, 500 + i * 18), line, font=F_XS, fill="#EAF2F8")
        d.text((112, 675), "OUTPUTS", font=F_XS, fill="#8CC63F")
        for i, metric in enumerate(metrics[:4]):
            rounded(d, (112, 702 + i * 34, 272, 728 + i * 34), "#EAF2F8", outline="#8CAAC0", width=1, radius=10)
            centered_text(d, (112, 702 + i * 34, 272, 728 + i * 34), metric, F_XS, fill=NAVY)

        # Main blueprint area
        d.text((340, 165), heading, font=F_H, fill=NAVY)
        d.text((340, 198), subtitle, font=F_S, fill="#435A6B")
        palette = [CYAN, GREEN, YELLOW, PURPLE, ORANGE, RED]
        row_tops = [250, 375, 500, 625, 750]
        node_no = 1
        previous_row_centers = []
        for r, row in enumerate(rows):
            y = row_tops[r]
            d.rounded_rectangle((335, y - 24, 1495, y + 92), radius=20, fill="#F7FAFC", outline="#C7D6E2", width=2)
            d.rectangle((335, y - 24, 347, y + 92), fill=palette[r % len(palette)])
            d.text((360, y - 16), row["label"], font=F_SB, fill=NAVY)
            cards = row["cards"]
            gap = 24
            card_w = int((1110 - gap * (len(cards) - 1)) / len(cards))
            x = 365
            centers = []
            for c, (title, detail) in enumerate(cards):
                box = (x, y + 10, x + card_w, y + 82)
                unified_node(d, box, node_no, title, detail, palette[(r + c) % len(palette)])
                centers.append(((box[0] + box[2]) // 2, box[1]))
                if c > 0:
                    arrow(d, (x - gap + 4, y + 46), (x - 4, y + 46), fill="#526D82", width=3)
                x += card_w + gap
                node_no += 1
            if previous_row_centers:
                for pc in previous_row_centers:
                    arrow(d, (pc[0], pc[1] + 92), (centers[min(len(centers) - 1, len(centers)//2)][0], y + 10), fill="#526D82", width=3)
            previous_row_centers = centers

        # Bottom status strip
        strip_y = 870
        d.rounded_rectangle((335, 850, 1495, 885), radius=12, fill="#EAF2F8", outline="#8CAAC0", width=2)
        d.text((355, 858), "Common decision layer: activity forecast -> confidence -> field action -> feedback update", font=F_S, fill=NAVY)
    return save_diagram(name, f"FIG. {fig_no} - {heading}", body)


def create_unified_figures():
    specs = [
        (
            "architecture.png",
            1,
            "Complete Prediction Architecture",
            "From field streams to pollination decision outputs.",
            [
                {"label": "INPUT STREAMS", "cards": [("Weather", "API and farm station readings"), ("Sensors", "Camera, acoustic and IoT nodes"), ("Crop", "Flowering stage and bloom density")]},
                {"label": "DATA CORE", "cards": [("Synchronize", "Match timestamp, field and crop"), ("Clean", "Missing values and noisy readings"), ("Feature Store", "Rolling windows and weather scores")]},
                {"label": "ENSEMBLE", "cards": [("Tree Models", "GBM and random forest"), ("Temporal Model", "Short-term activity trend"), ("Uncertainty", "Bayesian or quantile range")]},
                {"label": "DECISION", "cards": [("Calibrate", "Probability and confidence"), ("Risk Score", "Demand minus predicted activity"), ("Recommend", "Action timing and caution alert")]},
                {"label": "FEEDBACK", "cards": [("Observe", "Later visit counts"), ("Compare", "Forecast versus actual"), ("Update", "Weights and retraining queue")]},
            ],
            ["Index", "Risk", "Peak", "Alert"],
        ),
        (
            "prediction_flowchart.png",
            2,
            "Prediction Pipeline Flowchart",
            "One refresh cycle used for every field forecast.",
            [
                {"label": "START", "cards": [("Read Streams", "Weather, crop and activity data"), ("Validate", "Check range and timestamp"), ("Patch Gaps", "Fill or flag missing values")]},
                {"label": "FEATURES", "cards": [("Weather Lag", "Past 15-120 minute signal"), ("Stress Score", "Heat, wind and rain effect"), ("Bloom Demand", "Crop-stage pollination need")]},
                {"label": "PREDICT", "cards": [("Base Models", "Parallel activity estimates"), ("Meta Layer", "Weighted combined forecast"), ("Confidence", "Agreement and calibration")]},
                {"label": "DELIVER", "cards": [("Activity Index", "Field-level score"), ("Heat Map", "Location risk view"), ("Action Alert", "Spray, scout or hive timing")]},
                {"label": "STORE", "cards": [("Prediction Log", "Versioned output"), ("Field Result", "Observed visits"), ("Learning Queue", "Next update batch")]},
            ],
            ["Refresh", "Score", "Map", "Log"],
        ),
        (
            "training_loop.png",
            3,
            "Continuous Learning Loop",
            "How the model improves after new field observations.",
            [
                {"label": "COLLECT", "cards": [("Forecast Log", "Saved prediction history"), ("Actual Visits", "Manual or sensor counts"), ("Weather Context", "Conditions during outcome")]},
                {"label": "CHECK", "cards": [("Error", "Difference from actual"), ("Drift", "New weather pattern"), ("Quality", "Sensor health and gaps")]},
                {"label": "RETRAIN", "cards": [("Base Models", "Update selected learners"), ("Weights", "Shift ensemble importance"), ("Calibration", "Repair probability scale")]},
                {"label": "VALIDATE", "cards": [("Time Split", "Past predicts future"), ("Metrics", "MAE, AUC, Brier score"), ("Approval", "Deploy only if improved")]},
                {"label": "DEPLOY", "cards": [("Model Version", "Track release"), ("Dashboard", "Show new confidence"), ("Monitor", "Watch next cycle")]},
            ],
            ["MAE", "AUC", "Brier", "Drift"],
        ),
        (
            "dashboard_mockup.png",
            4,
            "Dashboard and Alert Interface",
            "The same forecast turned into farmer-readable outputs.",
            [
                {"label": "SUMMARY", "cards": [("Activity", "0.82 high"), ("Confidence", "91 percent"), ("Peak", "09:30-11:45")]},
                {"label": "FIELD VIEW", "cards": [("Heat Map", "Block-level activity"), ("Low Zones", "Pollination deficit areas"), ("Trend", "Hourly line forecast")]},
                {"label": "DRIVERS", "cards": [("Temperature", "Favorable range"), ("Wind", "Moderate disturbance"), ("Rain", "Recovery complete")]},
                {"label": "ACTION", "cards": [("Pesticide", "Avoid active window"), ("Scouting", "Check low zones"), ("Hive", "Move only if risk stays high")]},
                {"label": "EXPORT", "cards": [("Report", "PDF summary"), ("API", "Farm software link"), ("Archive", "Store field decision")]},
            ],
            ["0.82", "91%", "9-11", "Alert"],
        ),
        (
            "risk_matrix.png",
            5,
            "Pollination Deficit Risk Matrix",
            "How weather state and crop demand become action levels.",
            [
                {"label": "WEATHER", "cards": [("Calm Sunny", "Activity likely high"), ("Windy Cloudy", "Activity reduced"), ("Recent Rain", "Recovery delay")]},
                {"label": "CROP DEMAND", "cards": [("Early Bloom", "Demand rising"), ("Full Bloom", "Demand highest"), ("Late Bloom", "Demand falling")]},
                {"label": "PREDICTED ACTIVITY", "cards": [("High", "Sufficient visits"), ("Medium", "Watch conditions"), ("Low", "Deficit likely")]},
                {"label": "RISK", "cards": [("Low", "Normal monitoring"), ("Medium", "Scout and recheck"), ("High", "Support pollination")]},
                {"label": "FIELD ACTION", "cards": [("Spray Window", "Avoid activity peak"), ("Hive Plan", "Adjust placement"), ("Report", "Flag deficit zone")]},
            ],
            ["Low", "Med", "High", "Action"],
        ),
        (
            "model_stack.png",
            6,
            "Ensemble Model Stack",
            "Why the forecast does not depend on one algorithm.",
            [
                {"label": "FEATURES", "cards": [("Weather Vector", "Current and lagged values"), ("Crop Vector", "Stage and bloom density"), ("Activity Vector", "Recent visit signal")]},
                {"label": "BASE LEARNERS", "cards": [("GBM", "Nonlinear weather effects"), ("RF", "Noise-resistant baseline"), ("Temporal Net", "Short-term sequence pattern")]},
                {"label": "UNCERTAINTY", "cards": [("Quantile", "Upper and lower range"), ("Bayesian", "Confidence spread"), ("Classifier", "Active or inactive state")]},
                {"label": "STACKING", "cards": [("Meta Learner", "Combines estimates"), ("Calibration", "Corrects probability"), ("Disagreement", "Raises caution flag")]},
                {"label": "FINAL OUTPUT", "cards": [("Activity 0.84", "Score"), ("Low Risk", "Deficit result"), ("Peak 9-11 AM", "Timing recommendation")]},
            ],
            ["GBM", "RF", "TNN", "Meta"],
        ),
    ]
    return [create_unified_blueprint(*spec) for spec in specs]


def label_band(draw, text, y):
    draw.rounded_rectangle((95, y, 1505, y + 44), radius=12, fill="#EAF2F8", outline="#A9C0D1", width=2)
    draw.text((115, y + 10), text, font=F_SB, fill=NAVY)


def create_designed_architecture():
    def body(d):
        label_band(d, "Layered architecture: live sources -> feature intelligence -> ensemble forecast -> field decision", 150)
        layers = [
            ("INPUT STREAMS", 215, [("Weather API", CYAN), ("Farm station", CYAN), ("IoT sensors", GREEN), ("Camera count", GREEN), ("Crop stage", YELLOW), ("Field map", ORANGE)]),
            ("FEATURE INTELLIGENCE", 360, [("Sync", CYAN), ("Clean", GREEN), ("Lag windows", YELLOW), ("Stress scores", ORANGE), ("Feature store", PURPLE)]),
            ("ENSEMBLE CORE", 505, [("GBM", GREEN), ("Random forest", CYAN), ("Temporal net", PURPLE), ("Quantile model", ORANGE), ("Classifier", RED)]),
            ("DECISION OUTPUTS", 650, [("Activity index", GREEN), ("Peak window", YELLOW), ("Risk score", RED), ("Spray caution", ORANGE), ("Dashboard/API", PURPLE)]),
        ]
        for title, y, cards in layers:
            d.text((110, y + 16), title, font=F_SB, fill=NAVY)
            d.line((285, y + 27, 1460, y + 27), fill="#B7C9D8", width=3)
            x = 330
            for i, (txt, col) in enumerate(cards):
                rounded(d, (x, y, x + 165, y + 78), col, outline=NAVY, width=3, radius=18)
                centered_text(d, (x, y, x + 165, y + 78), txt, F_S, fill=NAVY)
                if i < len(cards) - 1:
                    arrow(d, (x + 165, y + 39), (x + 190, y + 39), fill="#536B7A", width=3)
                x += 190
        for x in [470, 660, 850, 1040, 1230]:
            arrow(d, (x, 293), (x, 360), fill="#536B7A", width=3)
            arrow(d, (x, 438), (x, 505), fill="#536B7A", width=3)
            arrow(d, (x, 583), (x, 650), fill="#536B7A", width=3)
        rounded(d, (375, 785, 1225, 850), "#102A43", outline="#102A43", width=2, radius=18)
        centered_text(d, (375, 785, 1225, 850), "Feedback loop: observed visits update calibration, weights, and future forecasts", F_H, fill="white")
        arrow(d, (1225, 817), (1365, 690), fill="#C00000", width=5)
        arrow(d, (375, 817), (240, 255), fill="#C00000", width=5)
    return save_diagram("architecture.png", "FIG. 1 - Complete Weather-Adaptive Prediction Architecture", body)


def create_designed_flow():
    def body(d):
        label_band(d, "Operational flowchart for each forecast refresh cycle", 150)
        steps = [
            ("01", "Collect streams", "weather, crop, field and pollinator data", CYAN),
            ("02", "Validate feed", "timestamp, range and missing-value checks", GREEN),
            ("03", "Build features", "lag, stress, recovery and bloom demand", YELLOW),
            ("04", "Run ensemble", "parallel model estimates", PURPLE),
            ("05", "Calibrate", "confidence and probability repair", ORANGE),
            ("06", "Issue outputs", "activity, risk, heat map and alert", RED),
            ("07", "Store feedback", "actual visits and model version", GREEN),
        ]
        x1, x2 = 230, 885
        y = 205
        for i, (num, title, detail, col) in enumerate(steps):
            x = x1 if i % 2 == 0 else x2
            rounded(d, (x, y, x + 485, y + 70), col, outline=NAVY, width=3, radius=20)
            d.ellipse((x + 18, y + 16, x + 60, y + 58), fill=NAVY)
            centered_text(d, (x + 18, y + 16, x + 60, y + 58), num, F_XS, fill="white")
            d.text((x + 78, y + 13), title, font=F_SB, fill=NAVY)
            d.text((x + 78, y + 40), detail, font=F_XS, fill="#263238")
            if i < len(steps) - 1:
                nx = x2 if i % 2 == 0 else x1
                ny = y + 82
                arrow(d, (x + 485, y + 35) if nx > x else (x, y + 35), (nx, ny + 35) if nx > x else (nx + 485, ny + 35), fill="#263238", width=4)
            y += 82
        d.rounded_rectangle((95, 820, 1505, 865), radius=14, fill="#EAF2F8", outline="#A9C0D1", width=2)
        centered_text(d, (95, 820, 1505, 865), "A refresh may run every 5, 15, 30, or 60 minutes based on sensor density and crop sensitivity.", F_S, fill=NAVY)
    return save_diagram("prediction_flowchart.png", "FIG. 2 - Prediction Pipeline Flowchart", body)


def create_designed_loop():
    def body(d):
        label_band(d, "Closed learning cycle for improving predictions after field feedback", 150)
        cx, cy = 800, 520
        items = [
            ("Forecast log", 800, 235, CYAN),
            ("Observed visits", 1160, 365, GREEN),
            ("Error analysis", 1160, 655, YELLOW),
            ("Retrain models", 800, 785, ORANGE),
            ("Deploy version", 440, 655, PURPLE),
            ("Live monitoring", 440, 365, RED),
        ]
        for i in range(len(items)):
            _, x, y, _ = items[i]
            _, nx, ny, _ = items[(i + 1) % len(items)]
            arrow(d, (x, y), (nx, ny), fill="#536B7A", width=5)
        for i, (txt, x, y, col) in enumerate(items, 1):
            rounded(d, (x - 145, y - 48, x + 145, y + 48), col, outline=NAVY, width=3, radius=22)
            centered_text(d, (x - 145, y - 48, x + 145, y + 48), f"{i}. {txt}", F_SB, fill=NAVY)
        d.ellipse((600, 390, 1000, 650), fill="#EAF2F8", outline=NAVY, width=5)
        centered_text(d, (640, 425, 960, 510), "Adaptive Core", F_H, fill=NAVY)
        centered_text(d, (660, 515, 940, 600), "weight update\ncalibration check\ndrift warning", F_S, fill="#263238")
        rounded(d, (610, 705, 990, 755), "#102A43", outline="#102A43", width=2, radius=16)
        centered_text(d, (610, 705, 990, 755), "Validation gate: deploy only when metrics improve", F_S, fill="white")
    return save_diagram("training_loop.png", "FIG. 3 - Continuous Learning and Recalibration Loop", body)


def create_designed_dashboard():
    def body(d):
        label_band(d, "Operator dashboard mockup: forecast, map, drivers, and action alert", 150)
        d.rounded_rectangle((100, 220, 1500, 835), radius=26, fill="#F8FBFD", outline=NAVY, width=4)
        d.rectangle((100, 220, 1500, 270), fill="#102A43")
        d.text((130, 235), "Pollinator Activity Console", font=F_H, fill="white")
        filters = ["Farm A", "Crop: Mustard", "Horizon: 6h", "Confidence: 91%"]
        for i, flt in enumerate(filters):
            rounded(d, (830 + i * 158, 231, 975 + i * 158, 260), "#EAF2F8", outline="#EAF2F8", width=1, radius=10)
            centered_text(d, (830 + i * 158, 231, 975 + i * 158, 260), flt, F_XS, fill=NAVY)
        cards = [("Activity Index", "0.82", GREEN), ("Peak Window", "09:30-11:45", YELLOW), ("Deficit Risk", "LOW", GREEN), ("Spray Caution", "AVOID PEAK", RED)]
        for i, (title, value, col) in enumerate(cards):
            x = 130 + i * 340
            rounded(d, (x, 300, x + 305, 410), col, outline=NAVY, width=3, radius=20)
            d.text((x + 20, 318), title, font=F_SB, fill=NAVY)
            centered_text(d, (x + 20, 345, x + 285, 398), value, font(28, True), fill=NAVY)
        rounded(d, (130, 455, 770, 800), "#FFFFFF", outline=NAVY, width=3, radius=18)
        d.text((160, 480), "Field Heat Map", font=F_SB, fill=NAVY)
        hm_cols = ["#D9EAD3", "#B6D7A8", "#93C47D", "#FFF2CC", "#F4CCCC"]
        for r in range(5):
            for c in range(8):
                fill = hm_cols[(r + c * 2) % len(hm_cols)]
                d.rounded_rectangle((165 + c * 68, 545 + r * 42, 220 + c * 68, 578 + r * 42), radius=8, fill=fill, outline="#777777")
        rounded(d, (825, 455, 1470, 800), "#FFFFFF", outline=NAVY, width=3, radius=18)
        d.text((855, 480), "Hourly Activity Trend", font=F_SB, fill=NAVY)
        base_x, base_y = 885, 740
        d.line((base_x, base_y, 1415, base_y), fill="#555555", width=3)
        d.line((base_x, 545, base_x, base_y), fill="#555555", width=3)
        pts = [(885, 710), (960, 675), (1035, 610), (1110, 575), (1185, 600), (1260, 665), (1335, 720), (1410, 735)]
        d.line(pts, fill=BLUE, width=8)
        for pnt in pts:
            d.ellipse((pnt[0]-8,pnt[1]-8,pnt[0]+8,pnt[1]+8), fill=BLUE)
    return save_diagram("dashboard_mockup.png", "FIG. 4 - Dashboard, Heat Map, and Alert Interface", body)


def create_designed_matrix():
    def body(d):
        label_band(d, "Decision matrix: converts weather state and crop demand into risk and action", 150)
        headers = ["Weather State", "Predicted Activity", "Crop Demand", "Risk", "Field Action"]
        rows = [
            ["Calm + sunny", "High", "High", "Low", "Normal monitoring"],
            ["Windy + cloudy", "Medium", "High", "Medium", "Scout / check hive"],
            ["Recent rain", "Low", "High", "High", "Delay spray / support"],
            ["Heat stress", "Low", "Medium", "Medium", "Use morning window"],
            ["Cold recovery", "Rising", "High", "Medium", "Recheck after warm-up"],
        ]
        x0, y0 = 120, 245
        widths = [275, 255, 245, 190, 365]
        x = x0
        for i, h in enumerate(headers):
            d.rectangle((x, y0, x + widths[i], y0 + 70), fill="#102A43", outline="white", width=2)
            centered_text(d, (x, y0, x + widths[i], y0 + 70), h, F_SB, fill="white")
            x += widths[i]
        for r, row in enumerate(rows):
            x = x0
            y = y0 + 70 + r * 88
            for c, cell in enumerate(row):
                fill = "#FFFFFF" if r % 2 == 0 else "#F7FAFC"
                if c == 3:
                    fill = {"Low": GREEN, "Medium": YELLOW, "High": RED}[cell]
                d.rectangle((x, y, x + widths[c], y + 88), fill=fill, outline="#8095A6", width=2)
                centered_text(d, (x, y, x + widths[c], y + 88), cell, F_S if c != 3 else F_SB, fill=NAVY)
                x += widths[c]
        rounded(d, (220, 790, 1380, 850), "#EAF2F8", outline=NAVY, width=3, radius=16)
        centered_text(d, (220, 790, 1380, 850), "Risk = crop pollination demand - weather-adjusted predicted activity + uncertainty penalty", F_H, fill=NAVY)
    return save_diagram("risk_matrix.png", "FIG. 5 - Pollination Deficit Risk Matrix", body)


def create_designed_stack():
    def body(d):
        label_band(d, "Model stack: varied learners feed one calibrated forecast", 150)
        levels = [
            ("Raw Inputs", 235, 6, [CYAN, GREEN, YELLOW, ORANGE, PURPLE, RED], ["Weather", "Crop", "Field", "Visits", "Sensors", "History"]),
            ("Feature Factory", 365, 5, [CYAN, GREEN, YELLOW, ORANGE, PURPLE], ["Lag", "Stress", "Bloom", "Trend", "Location"]),
            ("Base Learners", 495, 5, [GREEN, CYAN, PURPLE, ORANGE, RED], ["GBM", "RF", "TNN", "Quantile", "Classifier"]),
            ("Meta Layer", 625, 3, [CYAN, YELLOW, PURPLE], ["Stacking", "Calibration", "Disagreement"]),
            ("Forecast Output", 760, 3, [GREEN, YELLOW, RED], ["Activity 0.84", "Peak 9-11", "Risk Low"]),
        ]
        prev_centers = []
        for title, y, count, cols, labels in levels:
            d.text((130, y + 28), title, font=F_SB, fill=NAVY)
            total_w = 1040 - (count - 1) * 22
            box_w = total_w // count
            x = 390 + (5 - count) * 55
            centers = []
            for i in range(count):
                rounded(d, (x, y, x + box_w, y + 72), cols[i], outline=NAVY, width=3, radius=18)
                centered_text(d, (x, y, x + box_w, y + 72), labels[i], F_SB if count <= 3 else F_S, fill=NAVY)
                centers.append((x + box_w // 2, y))
                x += box_w + 22
            if prev_centers:
                target = centers[len(centers)//2]
                for pc in prev_centers:
                    arrow(d, (pc[0], pc[1] + 72), (target[0], y), fill="#536B7A", width=3)
            prev_centers = centers
        rounded(d, (1250, 265, 1478, 390), "#102A43", outline="#102A43", width=2, radius=18)
        centered_text(d, (1265, 285, 1465, 335), "Why ensemble?", F_SB, fill="white")
        centered_text(d, (1270, 335, 1460, 380), "Different learners catch different weather effects.", F_XS, fill="#EAF2F8")
        rounded(d, (1250, 520, 1478, 660), "#EAF2F8", outline=NAVY, width=3, radius=18)
        centered_text(d, (1265, 535, 1465, 640), "Meta layer keeps the final forecast stable when one model disagrees.", F_S, fill=NAVY)
    return save_diagram("model_stack.png", "FIG. 6 - Ensemble Model Stack and Confidence Engine", body)


def create_varied_designed_figures():
    return [
        create_designed_architecture(),
        create_designed_flow(),
        create_designed_loop(),
        create_designed_dashboard(),
        create_designed_matrix(),
        create_designed_stack(),
    ]


def clean_signature(src, name):
    img = Image.open(src).convert("RGB")
    # Only rotate obviously portrait captures. The new B signature is already horizontal.
    if img.height > img.width * 1.35:
        img = img.rotate(90, expand=True)
    gray = ImageOps.autocontrast(img.convert("L"))
    inv = ImageOps.invert(gray)
    bbox = inv.getbbox()
    if bbox:
        margin = 16
        img = img.crop((max(0, bbox[0] - margin), max(0, bbox[1] - margin), min(img.width, bbox[2] + margin), min(img.height, bbox[3] + margin)))
    img = ImageOps.autocontrast(img)
    max_w = 850
    if img.width > max_w:
        img = img.resize((max_w, int(img.height * max_w / img.width)))
    out = ASSETS / f"{name.lower().replace(' ', '_')}_signature.png"
    img.save(out)
    return out


def make_diagrams():
    return create_varied_designed_figures()


def docx_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def docx_para(doc, text="", label=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if label:
        r = p.add_run(label)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    return p


def style_docx_table(table, size=9):
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(size)


description = (
    "This invention describes a weather-aware prediction system for estimating pollinator activity in crop fields and ecological plots. Instead of treating pollination as a fixed seasonal event, the system reads changing weather conditions and combines them with crop stage, field context, and observed pollinator movement. Weather inputs may come from public weather feeds, local farm stations, low-cost IoT nodes, camera counters, acoustic sensors, and manual field logs. The collected data is arranged into short time windows, cleaned, and converted into practical indicators such as wind disturbance, rainfall recovery, heat stress, daylight phase, bloom demand, and recent visit trend. A group of machine learning models then works together to estimate likely pollinator activity, peak visiting periods, deficit risk, and prediction confidence. The output is delivered through a dashboard, heat map, alert panel, or API so that farmers and researchers can choose better times for pesticide spraying, scouting, hive placement, and other pollination-sensitive decisions."
)

problem = (
    "In many farms, pollinator activity is still judged by rough field observation or by general weather assumptions. That is not enough because a small change in wind, heat, rain, humidity, or cloud cover can shift pollinator movement within the same day. A field may look suitable in the morning and become poor for pollination by afternoon, or it may recover quickly after rain. Existing monitoring systems can count insects or watch hive movement, but they do not always tell the farmer what is likely to happen next. Single-model prediction tools also become weak when the weather pattern changes suddenly. Because of this, farmers may spray during active pollination, move hives at the wrong time, miss short activity windows, or notice pollination shortage only after the flowering stage has passed. The invention addresses this gap by giving a live, weather-adaptive prediction of pollinator activity with a confidence value and a practical recommendation."
)

objectives = [
    "To build a prediction framework that joins live weather data with crop stage, field layout, and pollinator observations.",
    "To use more than one machine learning model so that the system remains stable during changing weather.",
    "To generate an activity index, visitation probability, peak visiting window, deficit risk, and confidence value.",
    "To show the result through diagrams, heat maps, dashboard cards, and alerts that can guide field decisions.",
    "To improve the model over time by comparing forecasts with later field observations.",
]

prior_art = [
    ["1", "US20220125029A1", "Insect and creature monitoring system using detection devices and sensors.", "Detects pollinators but does not forecast activity from live weather streams.", "Transforms monitoring data into weather-adaptive forecasts with confidence scores."],
    ["2", "WO2020162926A1", "System for detecting numbers, health, and movement of pollinator populations.", "Monitoring oriented; lacks ensemble prediction and field-level decision logic.", "Uses monitored activity as feedback for ensemble training and drift control."],
    ["3", "US10064395B2", "Beehive monitoring system for beekeeper management.", "Hive-centered, not crop-field forecast oriented.", "Predicts crop-specific field activity and pollination deficit risk."],
    ["4", "US12075762B2", "Bee pollination monitoring using electronic devices and communication units.", "Tracks bee-related data without weather-adaptive activity window prediction.", "Forecasts peak windows and operational alerts using meteorological streams."],
    ["5", "WO2021060374A1", "Automatic pollination apparatus using imaging and pollination mechanisms.", "Hardware pollination solution rather than natural pollinator forecasting.", "Supports natural pollination management before artificial intervention."],
    ["6", "CN104255442A", "Artificial pollination device using airflow and pressure mechanisms.", "Mechanical pollen transfer only; no ecological prediction.", "Provides data-driven prediction and timing recommendations."],
    ["7", "WO2011090041A1", "Vibration pollination device for mechanical pollination assistance.", "Does not model weather-sensitive pollinator behavior.", "Determines when activity is sufficient or intervention is required."],
]

modules = [
    ("Real-Time Meteorological Stream Interface", "Collects temperature, humidity, wind, rainfall, radiation, cloud cover, dew point, pressure, and short-term forecasts."),
    ("Pollinator Observation Ingestion Module", "Accepts camera counts, acoustic events, visit logs, hive notes, trap readings, and manual survey entries."),
    ("Crop Phenology and Landscape Module", "Stores crop type, flowering stage, bloom density, field boundary, nearby habitat, and surrounding vegetation."),
    ("Feature Engineering Engine", "Builds lag values, moving averages, wind shock score, rainfall recovery time, heat stress score, daylight phase, and bloom demand."),
    ("Ensemble Learning Prediction Core", "Runs multiple models in parallel and combines their outputs instead of trusting one model alone."),
    ("Weather Adaptation and Drift Layer", "Checks whether the present weather pattern is different from earlier training conditions and updates the confidence value."),
    ("Explainability and Risk Module", "Shows which factors pushed the forecast up or down and whether pollination demand is likely to be unmet."),
    ("Dashboard and Alert Interface", "Displays activity cards, heat maps, timelines, pesticide caution windows, and recommended actions."),
]

claims = [
    "A computer implemented system for weather-adaptive pollinator activity prediction comprising a real-time meteorological stream interface, a pollinator observation ingestion module, a crop phenology module, a feature engineering engine, and an ensemble learning prediction core.",
    "The system as claimed in claim 1, wherein the meteorological stream interface receives at least temperature, humidity, rainfall, wind speed, solar radiation, cloud cover, dew point, pressure, and forecasted weather values.",
    "The system as claimed in claim 1, wherein the ensemble learning prediction core includes at least two models selected from gradient boosted decision trees, random forests, temporal neural networks, Bayesian models, quantile models, and logistic classifiers.",
    "The system as claimed in claim 1, wherein a stacking layer dynamically combines base model outputs according to weather regime, crop flowering stage, data completeness, and historical model performance.",
    "The system as claimed in claim 1, wherein the output comprises an activity index, visitation probability, expected activity count, peak activity window, pollination deficit risk, and confidence score.",
    "The system as claimed in claim 1, further comprising a drift detection layer configured to detect abnormal model disagreement or changing environmental regimes.",
    "The system as claimed in claim 1, further comprising an alert module configured to recommend pesticide avoidance windows, hive placement timing, scouting priority, or pollination support action.",
    "A method of predicting pollinator activity comprising receiving real-time weather data, receiving pollinator observations, generating rolling feature windows, processing the windows through a plurality of machine learning models, calibrating combined model predictions, and outputting a field-level activity forecast.",
    "The method as claimed in claim 8, wherein feedback from later observations is stored for retraining, recalibration, or dynamic model reweighting.",
]

advantages = [
    "Predicts future pollinator activity instead of only monitoring present activity.",
    "Handles sudden changes in wind, rain, heat, humidity, and solar radiation.",
    "Reduces pesticide exposure risk by identifying active pollination windows.",
    "Supports crop yield improvement through early deficit detection.",
    "Uses multiple models for stronger prediction under noisy field conditions.",
    "Provides explainable outputs for farmers, researchers, and policymakers.",
]

data_variables = [
    ("Meteorological Variables", "Temperature, relative humidity, wind speed, wind direction, wind gust, rainfall amount, rainfall duration, solar radiation, cloud cover, dew point, pressure, and short-term forecast values."),
    ("Crop Variables", "Crop species, cultivar, flowering stage, bloom density, flowering duration, row spacing, canopy structure, and expected pollination demand for the present crop stage."),
    ("Pollinator Variables", "Observed visit count, pollinator group, activity start time, peak movement time, flight interruption event, hive distance, and recent field-level visitation trend."),
    ("Spatial Variables", "Field boundary, block identifier, nearby natural habitat, vegetation strip, water source distance, pesticide buffer zone, and landscape fragmentation indicator."),
    ("Operational Variables", "Irrigation schedule, pesticide schedule, hive placement date, scouting time, crop protection action, and field feedback entered by the user."),
]

prototype_roadmap = [
    ("Month 1", "Prepare weather API connector, crop phenology input form, field database schema, and data cleaning pipeline."),
    ("Month 2", "Train baseline models using simulated or available observation datasets and construct rolling weather feature windows."),
    ("Month 3", "Build ensemble stacking layer, confidence score, dashboard prototype, heat map output, and pesticide caution alert logic."),
    ("Month 4", "Run field-style validation, compare predictions with observed visit counts, improve calibration, and prepare prototype report."),
]

validation_metrics = [
    ("Mean Absolute Error", "Compares predicted visit counts with observed pollinator visit counts."),
    ("AUC and F1-score", "Evaluate active versus low-activity classification under changing weather states."),
    ("Brier Score", "Checks whether predicted probabilities are properly calibrated."),
    ("Reliability Curve", "Compares confidence bands with actual forecast error."),
    ("Ablation Study", "Compares the full ensemble against weather-only, crop-only, and single-model baselines."),
]

commercialization_rows = [
    ("Precision Agriculture Platforms", "Integration as a pollination intelligence module inside smart farm dashboards."),
    ("Apiary Management Services", "Support for hive placement, movement timing, and pollination contract monitoring."),
    ("Crop Protection Advisory Systems", "Pesticide timing recommendations that reduce pollinator exposure during active windows."),
    ("Research Institutions", "Field-scale ecological monitoring for biodiversity, yield, and climate impact studies."),
    ("Agricultural Insurance Analytics", "Pollination deficit risk evidence for yield-risk modelling and advisory products."),
]


deployment_components = [
    ("Edge Data Collector", "A lightweight farm-side device or mobile gateway receives local sensor readings and forwards them to the prediction service."),
    ("Cloud Prediction Service", "A hosted service stores feature windows, executes ensemble models, and returns updated activity forecasts to user dashboards."),
    ("Farmer Dashboard", "A browser or mobile interface shows field heat maps, peak activity windows, confidence values, and pesticide caution periods."),
    ("Research Export Layer", "Researchers may export anonymized observation logs, forecast accuracy reports, and model validation summaries for field studies."),
]


user_workflow_steps = [
    ("Field Setup", "The user selects a farm, crop, field block, flowering stage, and optional pollinator category before activating the forecast run."),
    ("Stream Confirmation", "The dashboard checks whether weather feed, sensor input, crop stage, and recent observation records are available and marks missing streams."),
    ("Forecast Generation", "The ensemble prediction service computes activity index, peak activity window, confidence score, and pollination deficit risk."),
    ("Action Selection", "The user reviews pesticide caution windows, scouting priority zones, hive movement suggestions, and recommended recheck time."),
    ("Feedback Entry", "After the forecast window, the user or sensor system records actual visit counts so the model can compare expected and observed activity."),
]


use_disclosure_details = [
    "The disclosed invention is intended for academic IPR filing and future prototype development. It has not been launched as a commercial product, published as a public software release, or demonstrated as a market-ready service.",
    "Any experimental data used during prototype preparation may be generated from public weather feeds, controlled sample datasets, farm observations, or manually recorded field surveys. Such data will be used only to test prediction quality and not to limit the scope of the invention.",
    "The invention may be practiced as software, a cloud-based service, an edge-assisted farm device, or a dashboard module integrated into an existing farm management platform.",
    "The disclosure covers equivalent implementations where the same functional result is achieved using different machine learning libraries, different weather data providers, different farm sensor brands, or different dashboard technologies.",
    "The inventors reserve the right to extend the disclosure after prototype validation by adding experimental accuracy values, field photographs, dashboard screenshots, and comparative performance charts against baseline models.",
]


market_readiness_details = [
    ("Immediate Users", "Research farms, university field laboratories, apiary managers, and smart agriculture pilot projects can use the prototype for experimental forecasting and decision support."),
    ("Medium-Term Users", "Commercial orchards, seed production farms, oilseed producers, and managed pollination service providers can adopt the system after field validation."),
    ("Long-Term Users", "Government biodiversity programs, agricultural insurance platforms, and regional climate-resilience projects can use aggregated forecasts for monitoring and planning."),
    ("Subscription Model", "The system may be offered as a seasonal crop advisory subscription where users pay for field-level forecasts, alerts, and downloadable reports during flowering periods."),
    ("Licensing Model", "The prediction engine may be licensed to agri-tech companies that already operate farm dashboards and require a pollination intelligence module."),
    ("Pilot Deployment Model", "The first commercial trials may be conducted on high-value crops where pollination timing has measurable economic impact, such as orchard fruits, seed crops, oilseeds, and greenhouse crops."),
]


technology_stack_details = [
    ("Data Acquisition", "Weather APIs, farm weather stations, low-cost IoT sensor nodes, camera counters, acoustic event sensors, manual survey forms, and crop phenology records."),
    ("Processing Layer", "Time-series databases, feature engineering scripts, data validation routines, rolling window calculators, and missing-value treatment logic."),
    ("Prediction Layer", "Gradient boosting, random forest, temporal sequence models, quantile regression, calibration models, model disagreement checks, and stacking meta-learner."),
    ("Output Layer", "Dashboard cards, heat maps, field reports, alert notifications, downloadable summaries, and API integration for farm management systems."),
    ("Security and Access Layer", "User authentication, farm-level access control, data ownership rules, secure API keys for weather feeds, and audit logs for forecast decisions."),
    ("Maintenance Layer", "Model versioning, periodic retraining, sensor health checks, forecast error monitoring, and fallback rules when one data stream becomes unavailable."),
]


keyword_groups = [
    ("Core Technology Keywords", "Weather-adaptive forecasting, ensemble learning, pollinator activity prediction, real-time meteorological streams, calibrated confidence scoring."),
    ("Agriculture Keywords", "Precision agriculture, crop phenology, bloom demand, pollination deficit, hive placement, pesticide caution window, yield support."),
    ("Data Keywords", "IoT sensor stream, weather API, rolling feature window, field heat map, temporal validation, model drift detection."),
    ("Commercial Keywords", "Smart farming platform, apiary management, ecological monitoring, agricultural advisory, crop insurance analytics."),
    ("Research Keywords", "Field validation, pollination ecology, microclimate response, behavioral forecasting, biodiversity monitoring, crop-specific activity index."),
]


risk_mitigation_rows = [
    ("Missing Weather Feed", "Use last valid reading, nearby station fallback, and reduced confidence flag until the stream is restored."),
    ("Low Sensor Coverage", "Allow manual observation entry and generate field-level forecast with wider uncertainty band."),
    ("Unusual Weather Regime", "Trigger drift warning and increase reliance on models that performed better under similar historical conditions."),
    ("Crop Stage Error", "Request user confirmation of flowering stage and mark the prediction as provisional until corrected."),
    ("Model Disagreement", "Display a caution message and prevent the system from presenting a high-confidence recommendation."),
]


pilot_checklist_rows = [
    ("Weather Stream Test", "Confirm that local readings and external weather API values are synchronized within the accepted time tolerance."),
    ("Observation Test", "Compare camera, acoustic, and manual observation entries for the same field block during one selected flowering window."),
    ("Forecast Review", "Check whether the predicted peak activity period matches the observed peak visit interval within an acceptable margin."),
    ("Dashboard Review", "Verify that farmers can understand the activity score, risk category, and recommended action without technical assistance."),
    ("Retraining Review", "Store prediction results and actual field observations so that the next model update has labeled evidence."),
]


governance_rows = [
    ("Data Ownership", "Farm-level observations remain associated with the submitting user or institution and may be anonymized before research export."),
    ("Responsible Alerts", "Recommendations are advisory and should be combined with local agronomic judgment, pesticide labels, and pollinator safety rules."),
    ("Model Transparency", "The dashboard displays major weather and crop drivers so the user can understand why a forecast was produced."),
    ("Environmental Safety", "The system prioritizes avoidance of active pollination windows when suggesting pesticide timing or field interventions."),
    ("Audit Trail", "Forecasts, confidence values, and user decisions are stored so that later review can compare recommendation and outcome."),
    ("Regional Adaptation", "Models may be recalibrated by crop, region, and pollinator group to avoid assuming that all farms behave identically."),
]


def create_docx(diagrams, sigs):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)

    p = doc.add_paragraph("Annexure3b- Complete filing")
    p.runs[0].font.name = "Times New Roman"
    p = doc.add_paragraph("INVENTION DISCLOSURE FORM")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(16)
    docx_para(doc, "Details of Invention for better understanding:")
    docx_para(doc, TITLE, "1. TITLE: ")
    docx_heading(doc, "2. INTERNAL INVENTOR(S)/ STUDENT(S)", 2)
    for inv in inventors:
        table = doc.add_table(rows=6, cols=2)
        rows = [
            (f"{inv['label']}. Full name", inv["name"]),
            ("Mobile Number", inv["mobile"]),
            ("Email (personal)", inv["email"]),
            ("UID/Registration number", inv["reg"]),
            ("Address of Internal Inventors", inv["address"]),
            ("Signature", ""),
        ]
        for i, (k, v) in enumerate(rows):
            table.cell(i, 0).text = k
            table.cell(i, 1).text = v
        table.cell(5, 1).paragraphs[0].add_run().add_picture(str(sigs[inv["name"]]), width=Inches(2.15))
        style_docx_table(table)
        doc.add_paragraph()
    docx_para(doc, "None", "EXTERNAL INVENTOR(S): ")

    docx_heading(doc, "3. DESCRIPTION OF THE INVENTION", 2)
    docx_para(doc, description)
    docx_para(doc, problem, "A. PROBLEM ADDRESSED BY THE INVENTION: ")
    docx_heading(doc, "B. OBJECTIVE OF THE INVENTION", 2)
    for i, obj in enumerate(objectives, 1):
        docx_para(doc, obj, f"Objective {i}: ")

    docx_heading(doc, "C. STATE OF THE ART / RESEARCH GAP / NOVELTY", 2)
    table = doc.add_table(rows=1, cols=5)
    for i, h in enumerate(["Sr. No.", "Patent ID", "Abstract", "Research Gap", "Novelty"]):
        table.rows[0].cells[i].text = h
    for row in prior_art:
        cells = table.add_row().cells
        for i, item in enumerate(row):
            cells[i].text = item
    style_docx_table(table, 8)

    doc.add_section(WD_SECTION.NEW_PAGE)
    docx_heading(doc, "D. DETAILED DESCRIPTION", 2)
    docx_para(doc, "The proposed system is a modular AI platform in which field, weather, crop, and pollinator data move through a synchronized pipeline before being evaluated by multiple machine learning models. Each module is independently replaceable, enabling deployment on farms, research stations, greenhouses, orchards, apiary networks, and ecological monitoring sites.")
    docx_heading(doc, "I. Core Module Functions", 3)
    for name, body in modules:
        docx_para(doc, body, f"{name}: ")
    docx_heading(doc, "II. Step-by-Step Working Process", 3)
    for i, step in enumerate([
        "Collect weather, crop, landscape, and pollinator data from multiple sources.",
        "Synchronize all data streams by timestamp, field location, crop, and pollinator category.",
        "Create rolling weather features and crop-stage demand variables.",
        "Run multiple model families in parallel and compare their outputs.",
        "Combine outputs through stacking, calibration, and confidence scoring.",
        "Generate activity forecast, risk score, heat map, and operational alert.",
        "Store observed outcomes for retraining and continuous improvement.",
    ], 1):
        docx_para(doc, step, f"Step {i}: ")

    docx_heading(doc, "E. WORKING MODEL DIAGRAMS", 2)
    captions = [
        "Figure 1 - Complete Architecture",
        "Figure 2 - Prediction Pipeline",
        "Figure 3 - Learning Loop",
        "Figure 4 - Dashboard Interface",
        "Figure 5 - Risk Matrix",
        "Figure 6 - Ensemble Stack",
    ]
    figure_notes = [
        "The architecture view shows how raw sources are transformed into forecast outputs through feature engineering, ensemble modelling, calibration, and feedback.",
        "The prediction pipeline shows the operational path followed each time the system refreshes a field forecast.",
        "The learning loop shows how new observations return to the training store and help adjust model weights over time.",
        "The dashboard view shows how farmers or researchers would see activity index, field heat map, drivers, and caution alerts.",
        "The risk matrix links weather state and crop demand with a practical action so the output is usable in the field.",
        "The ensemble stack shows how separate models are combined into one calibrated forecast instead of depending on a single algorithm.",
    ]
    for idx, diagram in enumerate(diagrams):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        run = p.add_run(captions[idx])
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(diagram), width=Inches(7.0))
        docx_para(doc, figure_notes[idx])
        if idx != len(diagrams) - 1:
            doc.add_page_break()

    docx_heading(doc, "F. RESULTS AND ADVANTAGES / CLAIMS", 2)
    docx_heading(doc, "I. Claims", 3)
    for i, claim in enumerate(claims, 1):
        docx_para(doc, claim, f"{i}. ")
    docx_heading(doc, "II. Advantages", 3)
    for adv in advantages:
        doc.add_paragraph(adv, style="List Bullet")

    docx_heading(doc, "G. EXPANSION", 2)
    docx_para(doc, "The invention may be extended to multiple pollinator categories, including bees, butterflies, moths, flies, beetles, and birds. It may include additional variables such as soil moisture, pesticide exposure, floral diversity, hive health, pest pressure, disease pressure, and climate trend indicators. Equivalent implementations may use cloud computing, edge devices, mobile applications, farm management systems, IoT gateways, or research dashboards.")
    docx_heading(doc, "H. WORKING PROTOTYPE / FORMULATION / DESIGN / COMPOSITION", 2)
    docx_para(doc, "A conceptual design and complete architecture are ready. A prototype may be developed within 3 to 4 months using weather APIs, sample field observations, sensor datasets, Python-based machine learning libraries, and a dashboard interface.")
    docx_heading(doc, "I. EXISTING DATA", 2)
    table = doc.add_table(rows=1, cols=4)
    for i, h in enumerate(["S. No.", "Existing System", "Primary Function", "Key Limitation"]):
        table.rows[0].cells[i].text = h
    for i, row in enumerate([
        ["Manual field observation", "Counts pollinator visits", "Slow, costly, not predictive"],
        ["Static habitat models", "Estimate suitability", "Weak real-time weather response"],
        ["Sensor monitoring", "Detects present activity", "Limited forecasting"],
        ["Single ML model", "Predicts from selected variables", "Weak when weather changes quickly"],
        ["Artificial pollination devices", "Mechanically transfer pollen", "No natural activity intelligence"],
    ], 1):
        cells = table.add_row().cells
        cells[0].text = str(i)
        for j, item in enumerate(row):
            cells[j + 1].text = item
    style_docx_table(table, 8)

    doc.add_page_break()
    docx_heading(doc, "J. DATA VARIABLES AND SYSTEM INPUTS", 2)
    docx_para(doc, "The invention is designed to work with heterogeneous data because pollinator movement is affected by weather, crop stage, field design, and farm operations at the same time. The following variables form the preferred input structure for the prediction engine.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Variable Category"
    table.rows[0].cells[1].text = "Examples and Role in Prediction"
    for category, detail in data_variables:
        cells = table.add_row().cells
        cells[0].text = category
        cells[1].text = detail
    style_docx_table(table, 8)
    docx_para(doc, "These variables may be expanded with soil moisture, pesticide exposure, hive health indicators, flower resource diversity, vegetation index, or local biodiversity records without changing the basic architecture of the system.")

    doc.add_page_break()
    docx_heading(doc, "K. MODEL TRAINING, VALIDATION, AND RELIABILITY", 2)
    docx_para(doc, "The prediction engine is evaluated using time-based validation so that earlier observations are used to predict later activity windows. This avoids random data leakage and better represents the way the system will be used in real agricultural conditions.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Validation Method"
    table.rows[0].cells[1].text = "Purpose"
    for metric, detail in validation_metrics:
        cells = table.add_row().cells
        cells[0].text = metric
        cells[1].text = detail
    style_docx_table(table, 8)
    docx_para(doc, "Reliability is further improved by monitoring model disagreement. If different base learners produce conflicting outputs under unusual weather, the system lowers confidence and warns the user that the recommendation should be treated cautiously.")

    doc.add_page_break()
    docx_heading(doc, "L. PROTOTYPE ROADMAP AND COMMERCIAL SCOPE", 2)
    docx_para(doc, "A working prototype can be prepared in four development stages. The roadmap below shows how the conceptual architecture may be converted into a demonstrable academic prototype.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Timeline"
    table.rows[0].cells[1].text = "Planned Work"
    for month, detail in prototype_roadmap:
        cells = table.add_row().cells
        cells[0].text = month
        cells[1].text = detail
    style_docx_table(table, 8)
    docx_heading(doc, "Commercial Application Areas", 3)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Application Area"
    table.rows[0].cells[1].text = "Commercial Use"
    for area, detail in commercialization_rows:
        cells = table.add_row().cells
        cells[0].text = area
        cells[1].text = detail
    style_docx_table(table, 8)

    doc.add_page_break()
    docx_heading(doc, "M. DEPLOYMENT EXAMPLE AND USER WORKFLOW", 2)
    docx_para(doc, "In a practical farm deployment, the system may be operated as a hybrid edge-and-cloud service. A local gateway collects field readings and sends cleaned records to the prediction service. The cloud service performs model inference and returns an activity forecast to the user dashboard.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Deployment Component"
    table.rows[0].cells[1].text = "Function"
    for component, detail in deployment_components:
        cells = table.add_row().cells
        cells[0].text = component
        cells[1].text = detail
    style_docx_table(table, 8)
    docx_para(doc, "A typical user workflow begins with field selection, crop stage confirmation, and weather stream activation. The system then generates the activity forecast, marks caution periods, and stores the decision record for later validation.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Workflow Step"
    table.rows[0].cells[1].text = "Description"
    for step, detail in user_workflow_steps:
        cells = table.add_row().cells
        cells[0].text = step
        cells[1].text = detail
    style_docx_table(table, 8)

    doc.add_page_break()
    docx_heading(doc, "4. USE AND DISCLOSURE", 2)
    docx_para(doc, "The present invention has not been publicly disclosed or commercially used prior to this academic filing draft. This disclosure is made for academic and patent filing purposes.")
    for detail in use_disclosure_details:
        docx_para(doc, detail)

    docx_heading(doc, "5. POTENTIAL CHANCES OF COMMERCIALISATION: YES", 2)
    docx_para(doc, "The invention is commercially relevant to precision agriculture, pollination management, smart weather advisory systems, apiary services, ecological monitoring, agricultural analytics, and crop insurance risk analytics.")
    docx_para(doc, "The commercial value of the invention arises from the fact that pollination is directly connected to crop yield, fruit quality, seed set, and farm timing decisions, but most farm platforms still treat pollinator activity as an external uncertainty rather than a measurable forecast variable.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Commercial Stage"
    table.rows[0].cells[1].text = "Target Users and Readiness"
    for stage, detail in market_readiness_details:
        cells = table.add_row().cells
        cells[0].text = stage
        cells[1].text = detail
    style_docx_table(table, 8)
    docx_para(doc, "Commercial deployment may begin with a limited pilot in one crop and one region, followed by expansion to multiple crops once model calibration data has been collected. The system is especially suitable for crops where flowering periods are short and weather disruption can quickly reduce pollination success.")

    docx_heading(doc, "9. BASIC PATENTS / TECHNOLOGY USED", 2)
    docx_para(doc, "The invention may utilize weather APIs, IoT sensors, databases, dashboards, and machine learning frameworks in a novel integrated manner directed to weather-adaptive pollinator activity prediction.")
    docx_para(doc, "The use of standard technologies does not reduce the novelty of the invention because the inventive concept lies in the particular arrangement of these technologies for pollinator-specific activity forecasting, uncertainty scoring, and field action generation.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Technology Layer"
    table.rows[0].cells[1].text = "Example Components"
    for layer, detail in technology_stack_details:
        cells = table.add_row().cells
        cells[0].text = layer
        cells[1].text = detail
    style_docx_table(table, 8)
    docx_heading(doc, "10. FILING OPTIONS: PROVISIONAL", 2)
    docx_para(doc, "A provisional filing is recommended because the invention is at a conceptual and developmental stage and may later be supported with prototype results and field validation.")
    docx_para(doc, "The provisional filing route allows the inventors to secure an early priority date while continuing prototype development, field validation, dashboard refinement, and comparative model testing before a complete specification is prepared.")
    docx_para(doc, "During the provisional period, the inventors may add working screenshots, trained model summaries, field trial photographs, performance metrics, and expanded claims covering additional pollinator groups or crop categories.")
    docx_heading(doc, "11. KEYWORDS", 2)
    docx_para(doc, "Pollinator Activity Prediction, Ensemble Learning, Real-Time Meteorological Streams, Weather-Adaptive Forecasting, Precision Agriculture, Crop Phenology, Pollination Deficit Risk, Smart Farming, Ecological AI, Microclimate Analytics")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Keyword Group"
    table.rows[0].cells[1].text = "Keywords"
    for group, detail in keyword_groups:
        cells = table.add_row().cells
        cells[0].text = group
        cells[1].text = detail
    style_docx_table(table, 8)

    docx_heading(doc, "12. IMPLEMENTATION RISKS AND MITIGATION", 2)
    docx_para(doc, "Because agricultural fields are noisy real-world environments, the invention includes fallback behaviour for missing data, unusual weather, and uncertain crop-stage information. These safeguards prevent the system from presenting a misleading high-confidence forecast when input quality is weak.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Risk Condition"
    table.rows[0].cells[1].text = "Mitigation Strategy"
    for risk, detail in risk_mitigation_rows:
        cells = table.add_row().cells
        cells[0].text = risk
        cells[1].text = detail
    style_docx_table(table, 8)
    docx_para(doc, "The above safeguards are important because agricultural decisions may directly affect crop yield and pollinator safety. The system therefore avoids presenting a forecast as certain when data quality is incomplete or when environmental conditions differ from the available training history.")
    docx_heading(doc, "Pilot Evaluation Checklist", 3)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Pilot Check"
    table.rows[0].cells[1].text = "Evaluation Requirement"
    for check, detail in pilot_checklist_rows:
        cells = table.add_row().cells
        cells[0].text = check
        cells[1].text = detail
    style_docx_table(table, 8)
    docx_heading(doc, "13. DATA GOVERNANCE AND RESPONSIBLE USE", 2)
    docx_para(doc, "Because the invention processes farm observations, weather data, and pollinator activity records, the system is designed to support responsible data handling and transparent recommendations. The purpose of the forecast is to assist decision-making, not to replace agronomic judgment or local regulatory requirements.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Governance Aspect"
    table.rows[0].cells[1].text = "Implementation Approach"
    for aspect, detail in governance_rows:
        cells = table.add_row().cells
        cells[0].text = aspect
        cells[1].text = detail
    style_docx_table(table, 8)
    doc.add_page_break()
    p = doc.add_paragraph("NO OBJECTION CERTIFICATE")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(15)
    docx_para(doc, f'This is to certify that University/Organization Name or its associates shall have no objection if Lovely Professional University files an IPR (Patent/Copyright/Design/any other) entitled "{TITLE}" including the names of Aarav Kashyap Singh and Jai Dev Meena as inventors who are students studying in our University. Further Lovely Professional University shall not provide any financial assistance in respect of said IPR nor shall raise any objection later with respect to filing or commercialization of the said IPR or otherwise claim any right to the patent/invention at any stage.')
    doc.add_paragraph("\n\n(Authorised Signatory)")
    doc.save(DOCX_PATH)


def pdf_styles():
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("title", parent=base["Title"], fontName="Times-Bold", fontSize=15, leading=18, alignment=TA_CENTER, spaceAfter=9),
        "h1": ParagraphStyle("h1", parent=base["Heading1"], fontName="Times-Bold", fontSize=12, leading=15, spaceBefore=8, spaceAfter=5),
        "h2": ParagraphStyle("h2", parent=base["Heading2"], fontName="Times-Bold", fontSize=10.2, leading=12, spaceBefore=6, spaceAfter=4),
        "body": ParagraphStyle("body", parent=base["BodyText"], fontName="Times-Roman", fontSize=9.1, leading=11.5, alignment=TA_JUSTIFY, spaceAfter=4),
        "small": ParagraphStyle("small", parent=base["BodyText"], fontName="Times-Roman", fontSize=7.2, leading=8.6, alignment=TA_LEFT),
    }


def para(text, style):
    return Paragraph(text.replace("&", "&amp;"), style)


def rtable(rows, widths, size=7.2):
    st = pdf_styles()
    data = [[para(str(c), st["small"]) for c in row] for row in rows]
    t = Table(data, colWidths=widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#666666")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9EAF7")),
        ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("FONTSIZE", (0, 0), (-1, -1), size),
    ]))
    return t


def footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Times-Roman", 8)
    canvas.drawString(0.55 * inch, 0.38 * inch, "Weather-Adaptive Pollinator Activity Prediction - Visual Patent Draft")
    canvas.drawRightString(A4[0] - 0.55 * inch, 0.38 * inch, f"Page {doc.page}")
    canvas.restoreState()


def create_pdf(diagrams, sigs):
    st = pdf_styles()
    doc = SimpleDocTemplate(str(PDF_PATH), pagesize=A4, leftMargin=0.52 * inch, rightMargin=0.52 * inch, topMargin=0.54 * inch, bottomMargin=0.58 * inch)
    story = []
    story += [para("Annexure3b- Complete filing", st["body"]), para("INVENTION DISCLOSURE FORM", st["title"]), para("Details of Invention for better understanding:", st["body"]), para(f"<b>1. TITLE:</b> {TITLE}", st["body"]), para("<b>2. INTERNAL INVENTOR(S)/ STUDENT(S)</b>", st["h1"])]
    for inv in inventors:
        rows = [
            [f"{inv['label']}. Full name", inv["name"]],
            ["Mobile Number", inv["mobile"]],
            ["Email (personal)", inv["email"]],
            ["UID/Registration number", inv["reg"]],
            ["Address of Internal Inventors", inv["address"]],
            ["Signature", RLImage(str(sigs[inv["name"]]), width=1.95 * inch, height=0.55 * inch)],
        ]
        t = Table(rows, colWidths=[2.0 * inch, 4.55 * inch])
        t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.45, colors.grey), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("FONTSIZE", (0, 0), (-1, -1), 8.8), ("LEFTPADDING", (0, 0), (-1, -1), 4), ("RIGHTPADDING", (0, 0), (-1, -1), 4)]))
        story += [t, Spacer(1, 5)]
    story += [para("<b>EXTERNAL INVENTOR(S):</b> None", st["body"]), PageBreak(), para("3. DESCRIPTION OF THE INVENTION", st["h1"]), para(description, st["body"]), para(f"<b>A. PROBLEM ADDRESSED BY THE INVENTION:</b> {problem}", st["body"]), para("B. OBJECTIVE OF THE INVENTION", st["h1"])]
    for i, obj in enumerate(objectives, 1):
        story.append(para(f"<b>Objective {i}:</b> {obj}", st["body"]))
    story += [para("C. STATE OF THE ART / RESEARCH GAP / NOVELTY", st["h1"]), rtable([["Sr.", "Patent ID", "Abstract", "Research Gap", "Novelty"]] + prior_art, [0.34 * inch, 0.82 * inch, 1.68 * inch, 1.82 * inch, 1.86 * inch], 6.4)]

    story += [PageBreak(), para("D. DETAILED DESCRIPTION", st["h1"]), para("The proposed system is a modular AI platform in which field, weather, crop, and pollinator data move through a synchronized pipeline before being evaluated by multiple machine learning models. Each module is independently replaceable, enabling deployment on farms, research stations, greenhouses, orchards, apiary networks, and ecological monitoring sites.", st["body"]), para("I. Core Module Functions", st["h2"])]
    for name, body in modules:
        story.append(para(f"<b>{name}:</b> {body}", st["body"]))
    story.append(para("II. Step-by-Step Working Process", st["h2"]))
    for i, step in enumerate(["Collect weather, crop, landscape, and pollinator data from multiple sources.", "Synchronize all data streams by timestamp, field location, crop, and pollinator category.", "Create rolling weather features and crop-stage demand variables.", "Run multiple model families in parallel and compare outputs.", "Combine outputs through stacking, calibration, and confidence scoring.", "Generate activity forecast, risk score, heat map, and operational alert.", "Store observed outcomes for retraining and continuous improvement."], 1):
        story.append(para(f"<b>Step {i}:</b> {step}", st["body"]))

    story += [PageBreak(), para("E. WORKING MODEL DIAGRAMS", st["h1"])]
    captions = [
        "Figure 1 - Complete Architecture",
        "Figure 2 - Prediction Pipeline",
        "Figure 3 - Learning Loop",
        "Figure 4 - Dashboard Interface",
        "Figure 5 - Risk Matrix",
        "Figure 6 - Ensemble Stack",
    ]
    figure_notes = [
        "The architecture view shows how raw sources are transformed into forecast outputs through feature engineering, ensemble modelling, calibration, and feedback.",
        "The prediction pipeline shows the operational path followed each time the system refreshes a field forecast.",
        "The learning loop shows how new observations return to the training store and help adjust model weights over time.",
        "The dashboard view shows how farmers or researchers would see activity index, field heat map, drivers, and caution alerts.",
        "The risk matrix links weather state and crop demand with a practical action so the output is usable in the field.",
        "The ensemble stack shows how separate models are combined into one calibrated forecast instead of depending on a single algorithm.",
    ]
    for idx, diagram in enumerate(diagrams):
        block = [
            para(f"<b>{captions[idx]}</b>", st["h2"]),
            RLImage(str(diagram), width=7.05 * inch, height=4.40 * inch),
            Spacer(1, 5),
            para(figure_notes[idx], st["body"]),
        ]
        story.append(KeepTogether(block))
        if idx != len(diagrams) - 1:
            story.append(PageBreak())

    story += [PageBreak(), para("F. RESULTS AND ADVANTAGES / CLAIMS", st["h1"]), para("I. Claims", st["h2"])]
    for i, claim in enumerate(claims, 1):
        story.append(para(f"<b>{i}.</b> {claim}", st["body"]))
    story += [para("II. Advantages", st["h2"]), ListFlowable([ListItem(para(a, st["body"])) for a in advantages], bulletType="bullet", leftIndent=15), para("G. EXPANSION", st["h1"]), para("The invention may be extended to multiple pollinator categories, additional environmental variables, farm management systems, cloud services, edge devices, mobile dashboards, and large-scale ecological monitoring networks without changing the fundamental architecture.", st["body"]), para("H. WORKING PROTOTYPE / FORMULATION / DESIGN / COMPOSITION", st["h1"]), para("A conceptual design and complete architecture are ready. A prototype may be developed within 3 to 4 months using weather APIs, sample field observations, sensor datasets, Python-based machine learning libraries, and a dashboard interface.", st["body"]), para("I. EXISTING DATA", st["h1"])]
    story.append(rtable([["S. No.", "Existing System", "Primary Function", "Key Limitation"], ["1", "Manual field observation", "Counts pollinator visits", "Slow, costly, not predictive"], ["2", "Static habitat models", "Estimate suitability", "Weak real-time weather response"], ["3", "Sensor monitoring", "Detects present activity", "Limited forecasting"], ["4", "Single ML model", "Predicts from selected variables", "Weak when weather changes quickly"], ["5", "Artificial pollination devices", "Mechanically transfer pollen", "No natural activity intelligence"]], [0.45 * inch, 1.7 * inch, 2.15 * inch, 2.2 * inch], 7.2))
    story += [
        PageBreak(),
        para("J. DATA VARIABLES AND SYSTEM INPUTS", st["h1"]),
        para("The invention is designed to work with heterogeneous data because pollinator movement is affected by weather, crop stage, field design, and farm operations at the same time. The following variables form the preferred input structure for the prediction engine.", st["body"]),
        rtable([["Variable Category", "Examples and Role in Prediction"]] + [[a, b] for a, b in data_variables], [1.75 * inch, 4.95 * inch], 7.4),
        para("These variables may be expanded with soil moisture, pesticide exposure, hive health indicators, flower resource diversity, vegetation index, or local biodiversity records without changing the basic architecture of the system.", st["body"]),
        PageBreak(),
        para("K. MODEL TRAINING, VALIDATION, AND RELIABILITY", st["h1"]),
        para("The prediction engine is evaluated using time-based validation so that earlier observations are used to predict later activity windows. This avoids random data leakage and better represents the way the system will be used in real agricultural conditions.", st["body"]),
        rtable([["Validation Method", "Purpose"]] + [[a, b] for a, b in validation_metrics], [1.9 * inch, 4.8 * inch], 7.6),
        para("Reliability is further improved by monitoring model disagreement. If different base learners produce conflicting outputs under unusual weather, the system lowers confidence and warns the user that the recommendation should be treated cautiously.", st["body"]),
        PageBreak(),
        para("L. PROTOTYPE ROADMAP AND COMMERCIAL SCOPE", st["h1"]),
        para("A working prototype can be prepared in four development stages. The roadmap below shows how the conceptual architecture may be converted into a demonstrable academic prototype.", st["body"]),
        rtable([["Timeline", "Planned Work"]] + [[a, b] for a, b in prototype_roadmap], [1.2 * inch, 5.5 * inch], 7.7),
        para("Commercial Application Areas", st["h2"]),
        rtable([["Application Area", "Commercial Use"]] + [[a, b] for a, b in commercialization_rows], [2.0 * inch, 4.7 * inch], 7.4),
        PageBreak(),
        para("M. DEPLOYMENT EXAMPLE AND USER WORKFLOW", st["h1"]),
        para("In a practical farm deployment, the system may be operated as a hybrid edge-and-cloud service. A local gateway collects field readings and sends cleaned records to the prediction service. The cloud service performs model inference and returns an activity forecast to the user dashboard.", st["body"]),
        rtable([["Deployment Component", "Function"]] + [[a, b] for a, b in deployment_components], [2.0 * inch, 4.7 * inch], 7.6),
        para("A typical user workflow begins with field selection, crop stage confirmation, and weather stream activation. The system then generates the activity forecast, marks caution periods, and stores the decision record for later validation.", st["body"]),
        rtable([["Workflow Step", "Description"]] + [[a, b] for a, b in user_workflow_steps], [1.7 * inch, 5.0 * inch], 7.4),
    ]
    story += [
        PageBreak(),
        para("4. USE AND DISCLOSURE", st["h1"]),
        para("The present invention has not been publicly disclosed or commercially used prior to this academic filing draft. This disclosure is made for academic and patent filing purposes.", st["body"]),
    ]
    for detail in use_disclosure_details:
        story.append(para(detail, st["body"]))
    story += [
        para("5. POTENTIAL CHANCES OF COMMERCIALISATION: YES", st["h1"]),
        para("The invention is commercially relevant to precision agriculture, pollination management, smart weather advisory systems, apiary services, ecological monitoring, agricultural analytics, and crop insurance risk analytics.", st["body"]),
        para("The commercial value of the invention arises from the fact that pollination is directly connected to crop yield, fruit quality, seed set, and farm timing decisions, but most farm platforms still treat pollinator activity as an external uncertainty rather than a measurable forecast variable.", st["body"]),
        rtable([["Commercial Stage", "Target Users and Readiness"]] + [[a, b] for a, b in market_readiness_details], [1.8 * inch, 4.9 * inch], 7.4),
        para("Commercial deployment may begin with a limited pilot in one crop and one region, followed by expansion to multiple crops once model calibration data has been collected. The system is especially suitable for crops where flowering periods are short and weather disruption can quickly reduce pollination success.", st["body"]),
        para("9. BASIC PATENTS / TECHNOLOGY USED", st["h1"]),
        para("The invention may utilize weather APIs, IoT sensors, databases, dashboards, and machine learning frameworks in a novel integrated manner directed to weather-adaptive pollinator activity prediction.", st["body"]),
        para("The use of standard technologies does not reduce the novelty of the invention because the inventive concept lies in the particular arrangement of these technologies for pollinator-specific activity forecasting, uncertainty scoring, and field action generation.", st["body"]),
        rtable([["Technology Layer", "Example Components"]] + [[a, b] for a, b in technology_stack_details], [1.7 * inch, 5.0 * inch], 7.4),
        para("10. FILING OPTIONS: PROVISIONAL", st["h1"]),
        para("A provisional filing is recommended because the invention is at a conceptual and developmental stage and may later be supported with prototype results and field validation.", st["body"]),
        para("The provisional filing route allows the inventors to secure an early priority date while continuing prototype development, field validation, dashboard refinement, and comparative model testing before a complete specification is prepared.", st["body"]),
        para("During the provisional period, the inventors may add working screenshots, trained model summaries, field trial photographs, performance metrics, and expanded claims covering additional pollinator groups or crop categories.", st["body"]),
        para("11. KEYWORDS", st["h1"]),
        para("Pollinator Activity Prediction, Ensemble Learning, Real-Time Meteorological Streams, Weather-Adaptive Forecasting, Precision Agriculture, Crop Phenology, Pollination Deficit Risk, Smart Farming, Ecological AI, Microclimate Analytics", st["body"]),
        rtable([["Keyword Group", "Keywords"]] + [[a, b] for a, b in keyword_groups], [1.7 * inch, 5.0 * inch], 7.4),
        para("12. IMPLEMENTATION RISKS AND MITIGATION", st["h1"]),
        para("Because agricultural fields are noisy real-world environments, the invention includes fallback behaviour for missing data, unusual weather, and uncertain crop-stage information. These safeguards prevent the system from presenting a misleading high-confidence forecast when input quality is weak.", st["body"]),
        rtable([["Risk Condition", "Mitigation Strategy"]] + [[a, b] for a, b in risk_mitigation_rows], [1.9 * inch, 4.8 * inch], 7.4),
        para("The above safeguards are important because agricultural decisions may directly affect crop yield and pollinator safety. The system therefore avoids presenting a forecast as certain when data quality is incomplete or when environmental conditions differ from the available training history.", st["body"]),
        para("Pilot Evaluation Checklist", st["h2"]),
        rtable([["Pilot Check", "Evaluation Requirement"]] + [[a, b] for a, b in pilot_checklist_rows], [1.8 * inch, 4.9 * inch], 7.4),
        para("13. DATA GOVERNANCE AND RESPONSIBLE USE", st["h1"]),
        para("Because the invention processes farm observations, weather data, and pollinator activity records, the system is designed to support responsible data handling and transparent recommendations. The purpose of the forecast is to assist decision-making, not to replace agronomic judgment or local regulatory requirements.", st["body"]),
        rtable([["Governance Aspect", "Implementation Approach"]] + [[a, b] for a, b in governance_rows], [1.8 * inch, 4.9 * inch], 7.4),
    ]
    story += [PageBreak(), para("NO OBJECTION CERTIFICATE", st["title"]), para(f'This is to certify that University/Organization Name or its associates shall have no objection if Lovely Professional University files an IPR (Patent/Copyright/Design/any other) entitled "{TITLE}" including the names of Aarav Kashyap Singh and Jai Dev Meena as inventors who are students studying in our University. Further Lovely Professional University shall not provide any financial assistance in respect of said IPR nor shall raise any objection later with respect to filing or commercialization of the said IPR or otherwise claim any right to the patent/invention at any stage.', st["body"]), Spacer(1, 55), para("(Authorised Signatory)", st["body"])]
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


if __name__ == "__main__":
    diagrams = make_diagrams()
    sigs = {inv["name"]: clean_signature(inv["signature"], inv["name"]) for inv in inventors}
    create_docx(diagrams, sigs)
    create_pdf(diagrams, sigs)
    print(DOCX_PATH)
    print(PDF_PATH)
